# =============================================================================
# src/document_loaders/word_loader.py
# =============================================================================
"""
WordLoader Universal v2.5 - Extracción Completa de Documentos Word
==================================================================
Arquitectura Modular con Mixins para Funcionalidades Avanzadas

Features:
- ✅ Texto, tablas, imágenes, listas, formato, hyperlinks
- ✅ Gráficos embebidos (charts de Excel)
- ✅ Comentarios y notas al margen
- ✅ Footnotes/Endnotes
- ✅ Shapes y SmartArt
- ✅ Bookmarks internos
- ✅ Procesamiento paralelo para documentos grandes
- ✅ Metadata enriquecida para Qdrant
- ✅ Manejo robusto de errores

Autor: Sistema RAG Chatbot
Versión: 2.5
"""
import base64
import hashlib
import json
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from src.utils.date_utils import get_current_utc_iso
from src.utils.logger import get_logger
from .base_loader import BaseDocumentLoader, DocumentSection, ProcessedDocument

logger = get_logger(__name__)


# =============================================================================
# TEXT CLEANING UTILITIES
# =============================================================================

class TextCleaner:
    """Utility class for cleaning and optimizing text for RAG indexing"""

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean text by normalizing whitespace, removing control characters,
        and optimizing for embedding generation without losing context.

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text optimized for RAG
        """
        if not text:
            return ""

        # Remove control characters but preserve newlines and tabs
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)

        # Normalize whitespace: multiple spaces/tabs to single space
        text = re.sub(r'[ \t]+', ' ', text)

        # Normalize newlines: multiple consecutive newlines to double newline
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Remove excessive whitespace at line starts/ends
        text = re.sub(r'^\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s+$', '', text, flags=re.MULTILINE)

        # Remove empty lines at start/end of text
        text = text.strip()

        return text

    @staticmethod
    def clean_table_content(table_text: str) -> str:
        """
        Clean table content to be more compact while preserving structure.

        Args:
            table_text: Raw table text with | separators

        Returns:
            Cleaned table text
        """
        if not table_text:
            return ""

        # Clean each line
        lines = []
        for line in table_text.split('\n'):
            line = line.strip()
            if line:
                # Normalize spaces around | separators
                line = re.sub(r'\s*\|\s*', ' | ', line)
                lines.append(line)

        # Remove excessive empty lines
        cleaned = '\n'.join(lines)
        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)

        return cleaned.strip()

    @staticmethod
    def should_include_special_content(content_type: str, content: str) -> bool:
        """
        Determine if special content should be included in the final text.

        Args:
            content_type: Type of special content ('image', 'shape', 'comment', etc.)
            content: The content itself

        Returns:
            True if should include, False otherwise
        """
        # Exclude base64 images as they bloat payloads and aren't useful for text search
        if content_type == 'image':
            return False

        # Include comments and textboxes as they provide context
        if content_type in ['comment', 'textbox']:
            return len(content.strip()) > 10  # Only if meaningful content

        # Include shapes only if they have text
        if content_type == 'shape':
            return bool(content.strip())

        return True

    @staticmethod
    def format_special_content(content_type: str, content: str) -> str:
        """
        Format special content for inclusion in text.

        Args:
            content_type: Type of content
            content: The content

        Returns:
            Formatted content string
        """
        if not content.strip():
            return ""

        # Use more compact markers
        markers = {
            'textbox': '[TXT]',
            'comment': '[NOTE]',
            'shape': '[SHAPE]'
        }

        marker = markers.get(content_type, f'[{content_type.upper()}]')
        return f"{marker} {content.strip()}"


# =============================================================================
# MIXINS PARA FUNCIONALIDADES AVANZADAS
# =============================================================================

class ChartExtractorMixin:
    """Mixin para extracción de gráficos embebidos"""

    def _extract_charts(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extrae todos los gráficos embebidos del documento.

        Los gráficos de Word pueden ser:
        - Charts de Excel embebidos
        - Gráficos creados en Word

        Returns:
            List de dicts con información de gráficos
        """
        charts = []
        chart_index = 0

        try:
            # Buscar gráficos en relationships
            for rel in doc.part.rels.values():
                if "chart" in rel.target_ref.lower():
                    try:
                        chart_part = rel.target_part

                        # Extraer datos del gráfico
                        chart_info = {
                            'type': 'chart',
                            'index': chart_index,
                            'rel_id': rel.rId,
                            'target': rel.target_ref,
                            'content_type': chart_part.content_type if hasattr(chart_part, 'content_type') else 'unknown',
                        }

                        # Intentar extraer XML del gráfico para obtener datos
                        try:
                            chart_xml = chart_part.blob.decode('utf-8')
                            chart_info['has_data'] = True
                            chart_info['xml_preview'] = chart_xml[:200] + '...'
                        except:
                            chart_info['has_data'] = False

                        charts.append(chart_info)
                        chart_index += 1

                        if hasattr(self, 'stats'):
                            self.stats['charts'] = self.stats.get('charts', 0) + 1

                    except Exception as e:
                        logger.debug(f"Error extracting chart {rel.rId}: {e}")

        except Exception as e:
            logger.debug(f"Error in chart extraction: {e}")

        return charts

    def _extract_chart_from_paragraph(self, para: Paragraph) -> Optional[Dict[str, Any]]:
        """
        Extrae gráfico específico de un párrafo.

        Returns:
            Dict con info del gráfico o None
        """
        try:
            # Buscar elementos de gráfico en el XML del párrafo
            chart_elements = para._element.xpath('.//c:chart')

            if chart_elements:
                chart_elem = chart_elements[0]
                rId = chart_elem.get(qn('r:id'))

                if rId:
                    try:
                        chart_part = para.part.rels[rId].target_part

                        return {
                            'type': 'chart',
                            'rel_id': rId,
                            'content_type': chart_part.content_type if hasattr(chart_part, 'content_type') else 'unknown',
                            'location': 'inline',
                        }
                    except:
                        pass

        except Exception as e:
            logger.debug(f"Error extracting chart from paragraph: {e}")

        return None


class CommentExtractorMixin:
    """Mixin para extracción de comentarios y notas al margen"""

    def _extract_all_comments(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extrae todos los comentarios del documento.

        Returns:
            List de dicts con información de comentarios
        """
        comments = []

        try:
            # Buscar parte de comentarios
            for rel in doc.part.rels.values():
                if "comments" in rel.target_ref.lower():
                    try:
                        comments_part = rel.target_part

                        # Parsear XML de comentarios
                        comments_xml = comments_part.element

                        for comment in comments_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment'):
                            comment_id = comment.get(qn('w:id'))
                            author = comment.get(qn('w:author'), 'Unknown')
                            date = comment.get(qn('w:date'), '')

                            # Extraer texto del comentario
                            text_elements = comment.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            comment_text = ''.join([t.text for t in text_elements if t.text])

                            if comment_text:
                                comments.append({
                                    'type': 'comment',
                                    'id': comment_id,
                                    'author': author,
                                    'date': date,
                                    'text': comment_text.strip(),
                                })

                                if hasattr(self, 'stats'):
                                    self.stats['comments'] = self.stats.get('comments', 0) + 1

                    except Exception as e:
                        logger.debug(f"Error parsing comments part: {e}")

        except Exception as e:
            logger.debug(f"Error in comment extraction: {e}")

        return comments

    def _extract_paragraph_comments(self, para: Paragraph, all_comments: List[Dict]) -> List[Dict[str, Any]]:
        """
        Extrae comentarios asociados a un párrafo específico.

        Args:
            para: Párrafo a analizar
            all_comments: Lista de todos los comentarios del documento

        Returns:
            List de comentarios relacionados
        """
        paragraph_comments = []

        try:
            # Buscar referencias a comentarios en el párrafo
            comment_refs = para._element.xpath('.//w:commentReference')

            for ref in comment_refs:
                comment_id = ref.get(qn('w:id'))

                # Buscar el comentario correspondiente
                for comment in all_comments:
                    if comment.get('id') == comment_id:
                        paragraph_comments.append(comment)
                        break

        except Exception as e:
            logger.debug(f"Error extracting paragraph comments: {e}")

        return paragraph_comments


class FootnoteExtractorMixin:
    """Mixin para extracción de footnotes y endnotes"""

    def _extract_footnotes(self, doc: Document) -> Tuple[List[Dict], List[Dict]]:
        """
        Extrae footnotes y endnotes del documento.

        Returns:
            Tuple (footnotes_list, endnotes_list)
        """
        footnotes = []
        endnotes = []

        try:
            # Buscar parte de footnotes
            for rel in doc.part.rels.values():
                target = rel.target_ref.lower()

                if "footnote" in target and "endnote" not in target:
                    footnotes.extend(self._parse_notes_part(rel.target_part, 'footnote'))
                elif "endnote" in target:
                    endnotes.extend(self._parse_notes_part(rel.target_part, 'endnote'))

        except Exception as e:
            logger.debug(f"Error in footnote/endnote extraction: {e}")

        if hasattr(self, 'stats'):
            self.stats['footnotes'] = len(footnotes)
            self.stats['endnotes'] = len(endnotes)

        return footnotes, endnotes

    def _parse_notes_part(self, notes_part, note_type: str) -> List[Dict[str, Any]]:
        """
        Parsea una parte de notas (footnotes o endnotes).

        Args:
            notes_part: Parte del documento con notas
            note_type: 'footnote' o 'endnote'

        Returns:
            List de dicts con información de notas
        """
        notes = []

        try:
            notes_xml = notes_part.element
            note_tag = f'{{{notes_xml.nsmap["w"]}}}{note_type}'

            for note in notes_xml.findall(f'.//{note_tag}'):
                note_id = note.get(qn('w:id'))

                # Extraer texto de la nota
                text_elements = note.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                note_text = ''.join([t.text for t in text_elements if t.text])

                if note_text and note_id not in ['-1', '0']:  # Saltar notas de separación
                    notes.append({
                        'type': note_type,
                        'id': note_id,
                        'text': note_text.strip(),
                    })

        except Exception as e:
            logger.debug(f"Error parsing {note_type}s: {e}")

        return notes

    def _extract_paragraph_notes(self, para: Paragraph) -> Dict[str, List[str]]:
        """
        Extrae referencias a footnotes/endnotes en un párrafo.

        Returns:
            Dict con 'footnote_refs' y 'endnote_refs'
        """
        note_refs = {
            'footnote_refs': [],
            'endnote_refs': []
        }

        try:
            # Buscar referencias a footnotes
            footnote_refs = para._element.xpath('.//w:footnoteReference')
            for ref in footnote_refs:
                ref_id = ref.get(qn('w:id'))
                if ref_id:
                    note_refs['footnote_refs'].append(ref_id)

            # Buscar referencias a endnotes
            endnote_refs = para._element.xpath('.//w:endnoteReference')
            for ref in endnote_refs:
                ref_id = ref.get(qn('w:id'))
                if ref_id:
                    note_refs['endnote_refs'].append(ref_id)

        except Exception as e:
            logger.debug(f"Error extracting paragraph note references: {e}")

        return note_refs


class ShapeExtractorMixin:
    """Mixin para extracción de shapes y SmartArt"""

    def _extract_shapes(self, para: Paragraph) -> List[Dict[str, Any]]:
        """
        Extrae shapes (formas) de un párrafo.

        Incluye:
        - Shapes básicas (rectángulos, círculos, etc.)
        - WordArt
        - Diagramas

        Returns:
            List de dicts con información de shapes
        """
        shapes = []

        try:
            # Buscar elementos de shape en el XML
            shape_elements = para._element.xpath('.//v:shape | .//wps:wsp')

            for idx, shape in enumerate(shape_elements):
                shape_info = {
                    'type': 'shape',
                    'index': idx,
                }

                # Intentar extraer texto del shape (si tiene)
                text_elements = shape.xpath('.//w:t')
                shape_text = ''.join([t.text for t in text_elements if t.text])

                if shape_text:
                    shape_info['text'] = shape_text.strip()
                    shape_info['has_text'] = True
                else:
                    shape_info['has_text'] = False

                # Intentar extraer tipo de shape
                try:
                    shape_type = shape.get('type', 'unknown')
                    shape_info['shape_type'] = shape_type
                except:
                    shape_info['shape_type'] = 'unknown'

                shapes.append(shape_info)

                if hasattr(self, 'stats'):
                    self.stats['shapes'] = self.stats.get('shapes', 0) + 1

        except Exception as e:
            logger.debug(f"Error extracting shapes: {e}")

        return shapes

    def _extract_smartart(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extrae diagramas SmartArt del documento.

        SmartArt es más complejo ya que tiene estructura jerárquica.

        Returns:
            List de dicts con información de SmartArt
        """
        smartart_list = []

        try:
            # Buscar partes de SmartArt en relationships
            for rel in doc.part.rels.values():
                if "diagramData" in rel.target_ref.lower():
                    try:
                        smartart_part = rel.target_part

                        smartart_info = {
                            'type': 'smartart',
                            'rel_id': rel.rId,
                            'target': rel.target_ref,
                        }

                        # Intentar extraer texto del SmartArt
                        try:
                            smartart_xml = smartart_part.blob.decode('utf-8')

                            # Buscar elementos de texto en el XML
                            # (esto es simplificado, el XML de SmartArt es complejo)
                            text_matches = re.findall(r'<[^>]*:t[^>]*>([^<]+)</[^>]*:t>', smartart_xml)
                            if text_matches:
                                smartart_info['texts'] = text_matches
                                smartart_info['has_text'] = True
                            else:
                                smartart_info['has_text'] = False

                        except:
                            smartart_info['has_text'] = False

                        smartart_list.append(smartart_info)

                        if hasattr(self, 'stats'):
                            self.stats['smartart'] = self.stats.get('smartart', 0) + 1

                    except Exception as e:
                        logger.debug(f"Error extracting SmartArt {rel.rId}: {e}")

        except Exception as e:
            logger.debug(f"Error in SmartArt extraction: {e}")

        return smartart_list


class BookmarkExtractorMixin:
    """Mixin para extracción de bookmarks internos"""

    def _extract_all_bookmarks(self, doc: Document) -> Dict[str, Dict[str, Any]]:
        """
        Extrae todos los bookmarks del documento.

        Los bookmarks son marcadores internos que permiten navegación
        y referencias cruzadas.

        Returns:
            Dict con bookmark_name -> info
        """
        bookmarks = {}

        try:
            # Buscar todos los bookmarks en el documento
            for para in doc.paragraphs:
                para_bookmarks = self._extract_paragraph_bookmarks(para)
                bookmarks.update(para_bookmarks)

        except Exception as e:
            logger.debug(f"Error in bookmark extraction: {e}")

        if hasattr(self, 'stats'):
            self.stats['bookmarks'] = len(bookmarks)

        return bookmarks

    def _extract_paragraph_bookmarks(self, para: Paragraph) -> Dict[str, Dict[str, Any]]:
        """
        Extrae bookmarks de un párrafo específico.

        Returns:
            Dict con bookmark_name -> info
        """
        bookmarks = {}

        try:
            # Buscar elementos de bookmark start
            bookmark_starts = para._element.xpath('.//w:bookmarkStart')

            for bookmark in bookmark_starts:
                bookmark_name = bookmark.get(qn('w:name'))
                bookmark_id = bookmark.get(qn('w:id'))

                if bookmark_name:
                    bookmarks[bookmark_name] = {
                        'type': 'bookmark',
                        'id': bookmark_id,
                        'name': bookmark_name,
                        'paragraph_text': para.text[:100] + '...' if len(para.text) > 100 else para.text,
                    }

        except Exception as e:
            logger.debug(f"Error extracting paragraph bookmarks: {e}")

        return bookmarks


class ParallelProcessingMixin:
    """Mixin para procesamiento paralelo de documentos grandes"""

    def _should_use_parallel_processing(self, doc: Document) -> bool:
        """
        Determina si el documento es lo suficientemente grande
        para beneficiarse del procesamiento paralelo.

        Args:
            doc: Documento a analizar

        Returns:
            bool: True si debe usar procesamiento paralelo
        """
        # Umbral: documentos con más de 100 párrafos o 50 tablas
        return len(doc.paragraphs) > 100 or len(doc.tables) > 50

    def _process_sections_parallel(
        self,
        sections: List[DocumentSection],
        max_workers: int = 4
    ) -> List[DocumentSection]:
        """
        Procesa secciones en paralelo para mejorar rendimiento.

        Útil para documentos muy grandes donde cada sección puede
        procesarse independientemente.

        Args:
            sections: Lista de secciones a procesar
            max_workers: Número máximo de workers paralelos

        Returns:
            Lista de secciones procesadas
        """
        if not sections or len(sections) < 10:
            # No vale la pena paralelizar para documentos pequeños
            return sections

        processed_sections = []

        try:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Enviar cada sección a procesamiento
                future_to_section = {
                    executor.submit(self._enrich_section_metadata, section): section
                    for section in sections
                }

                # Recolectar resultados
                for future in as_completed(future_to_section):
                    try:
                        enriched_section = future.result()
                        processed_sections.append(enriched_section)
                    except Exception as e:
                        original_section = future_to_section[future]
                        logger.error(f"Error processing section '{original_section.title}': {e}")
                        processed_sections.append(original_section)

        except Exception as e:
            logger.error(f"Error in parallel processing: {e}")
            return sections  # Fallback a secciones originales

        # Reordenar para mantener el orden original
        section_order = {id(s): i for i, s in enumerate(sections)}
        processed_sections.sort(key=lambda s: section_order.get(id(s), 999999))

        return processed_sections

    def _enrich_section_metadata(self, section: DocumentSection) -> DocumentSection:
        """
        Enriquece la metadata de una sección con análisis adicional.

        Esta función está diseñada para ser ejecutada en paralelo.

        Args:
            section: Sección a enriquecer

        Returns:
            Sección con metadata enriquecida
        """
        try:
            # Análisis adicional de la sección
            content = section.content

            # Análisis de complejidad
            section.metadata['enrichment'] = {
                'char_count': len(content),
                'sentence_count': len(re.findall(r'[.!?]+', content)),
                'avg_word_length': sum(len(word) for word in content.split()) / max(len(content.split()), 1),
                'has_technical_terms': bool(re.search(r'\b[A-Z]{2,}\b', content)),
                'code_blocks': len(re.findall(r'```[\s\S]*?```', content)),
            }

        except Exception as e:
            logger.debug(f"Error enriching section metadata: {e}")

        return section


# =============================================================================
# CLASE PRINCIPAL CON TODOS LOS MIXINS
# =============================================================================

class WordLoaderUniversal(
    ChartExtractorMixin,
    CommentExtractorMixin,
    FootnoteExtractorMixin,
    ShapeExtractorMixin,
    BookmarkExtractorMixin,
    ParallelProcessingMixin,
    BaseDocumentLoader
):
    """
    Cargador Universal de Documentos Word con Extracción Completa

    Hereda de múltiples mixins para funcionalidades avanzadas:
    - ChartExtractorMixin: Gráficos embebidos
    - CommentExtractorMixin: Comentarios y notas
    - FootnoteExtractorMixin: Footnotes y endnotes
    - ShapeExtractorMixin: Shapes y SmartArt
    - BookmarkExtractorMixin: Bookmarks internos
    - ParallelProcessingMixin: Procesamiento paralelo

    Características:
    - Extrae texto, tablas, imágenes, listas, formato
    - Detecta hyperlinks, ecuaciones, cuadros de texto
    - Metadata enriquecida para cada chunk
    - Compatible con Qdrant via FileProcessor
    - Manejo robusto de errores
    - Procesamiento paralelo para documentos grandes
    """

    def __init__(self, enable_parallel: bool = True, max_workers: int = 4, clean_text: bool = True):
        """
        Inicializa el loader con opciones configurables.

        Args:
            enable_parallel: Habilitar procesamiento paralelo
            max_workers: Número máximo de workers para procesamiento paralelo
            clean_text: Habilitar limpieza de texto para optimizar RAG
        """
        super().__init__()
        self.supported_extensions = {'.docx'}
        self.enable_parallel = enable_parallel
        self.max_workers = max_workers
        self.clean_text = clean_text
        self.text_cleaner = TextCleaner() if clean_text else None

        # Estadísticas de extracción extendidas
        self.stats = {
            "paragraphs": 0,
            "tables": 0,
            "nested_tables": 0,  # NUEVO: Tablas anidadas
            "images": 0,
            "hyperlinks": 0,
            "lists": 0,
            "textboxes": 0,
            "equations": 0,
            "charts": 0,
            "comments": 0,
            "footnotes": 0,
            "endnotes": 0,
            "shapes": 0,
            "smartart": 0,
            "bookmarks": 0,
        }

        # Cache para elementos extraídos globalmente
        self._global_cache = {
            'comments': [],
            'footnotes': [],
            'endnotes': [],
            'charts': [],
            'smartart': [],
            'bookmarks': {},
        }

    def load(self, file_path: Path, original_filename: str = None) -> ProcessedDocument:
        """
        Carga un documento Word con extracción completa.

        Compatible con FileProcessor.process_file()

        Args:
            file_path: Ruta al archivo .docx
            original_filename: Nombre original del archivo (opcional)

        Returns:
            ProcessedDocument con metadata completa para Qdrant
        """
        logger.info(f"🚀 Loading Word document: {file_path}")

        # Reset estadísticas y cache
        self.stats = {k: 0 for k in self.stats}
        self._global_cache = {
            'comments': [],
            'footnotes': [],
            'endnotes': [],
            'charts': [],
            'smartart': [],
            'bookmarks': {},
        }

        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"❌ Error loading document: {e}")
            raise ValueError(f"No se pudo cargar el documento: {e}")

        # 1. Extraer elementos globales (una sola vez)
        logger.info("📦 Extracting global elements...")
        self._extract_global_elements(doc)

        # 2. Extraer metadata del documento
        logger.info("📦 Extracting metadata...")
        metadata = self._extract_metadata(doc, file_path)

        # Agregar elementos globales a metadata
        metadata['global_elements'] = {
            'chart_count': len(self._global_cache['charts']),
            'comment_count': len(self._global_cache['comments']),
            'footnote_count': len(self._global_cache['footnotes']),
            'endnote_count': len(self._global_cache['endnotes']),
            'smartart_count': len(self._global_cache['smartart']),
            'bookmark_count': len(self._global_cache['bookmarks']),
        }

        # 3. Extraer secciones con metadata enriquecida
        logger.info("📦 Extracting sections...")
        sections = self.extract_sections(doc, metadata)

        # 4. Procesamiento paralelo si está habilitado y el documento es grande
        if self.enable_parallel and self._should_use_parallel_processing(doc):
            logger.info(f"⚡ Using parallel processing with {self.max_workers} workers")
            sections = self._process_sections_parallel(sections, self.max_workers)

        logger.info(f"✅ Extracted {len(sections)} sections from {file_path}")
        logger.info(f"📊 Stats: {json.dumps(self.stats, indent=2)}")

        # 5. Generar contenido completo
        logger.info("📦 Generating full content...")
        content = self._generate_full_content(sections)

        # 6. Convertir a ruta relativa
        logger.info("📦 Converting to relative path...")
        abs_path = file_path if file_path.is_absolute() else file_path.resolve()
        try:
            relative_path = abs_path.relative_to(Path.cwd())
        except ValueError:
            relative_path = abs_path

        return ProcessedDocument(
            file_path=str(relative_path),
            file_name=file_path.name,
            original_filename=original_filename or file_path.name,
            content=content,
            sections=sections,
            metadata=metadata
        )

    def _extract_global_elements(self, doc: Document):
        """
        Extrae elementos globales del documento una sola vez.

        Esto mejora el rendimiento al evitar múltiples pasadas.
        """
        # Extraer comentarios
        self._global_cache['comments'] = self._extract_all_comments(doc)

        # Extraer footnotes y endnotes
        footnotes, endnotes = self._extract_footnotes(doc)
        self._global_cache['footnotes'] = footnotes
        self._global_cache['endnotes'] = endnotes

        # Extraer gráficos
        self._global_cache['charts'] = self._extract_charts(doc)

        # Extraer SmartArt
        self._global_cache['smartart'] = self._extract_smartart(doc)

        # Extraer bookmarks
        self._global_cache['bookmarks'] = self._extract_all_bookmarks(doc)

    # =========================================================================
    # EXTRACCIÓN DE METADATA
    # =========================================================================

    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """
        Extrae metadata completa del documento Word.

        Esta metadata se agrega a CADA chunk en Qdrant vía FileProcessor.

        Returns:
            dict: Metadatos del documento para indexar en Qdrant
        """
        core_props = doc.core_properties

        # Estadísticas del archivo
        logger.info("📦 Extracting file stats...")
        file_stats = file_path.stat()

        # Contar elementos
        logger.info("📦 Counting elements...")
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)

        # Detectar idioma (simple heurística basada en contenido)
        sample_text = " ".join([p.text for p in doc.paragraphs[:10] if p.text.strip()])
        detected_language = self._detect_language(sample_text)
        logger.info(f"📦 Detected language: {detected_language}")

        # Función helper para extraer propiedades de forma segura
        def safe_get_prop(prop_name: str, default: str = '') -> str:
            """Extrae propiedad de forma segura, retorna default si no existe"""
            try:
                value = getattr(core_props, prop_name, None)
                return str(value) if value is not None else default
            except (AttributeError, TypeError):
                return default

        # Extraer fechas de forma segura
        def safe_get_date(prop_name: str) -> str:
            """Extrae fecha de forma segura, retorna ISO string o vacío"""
            try:
                date_obj = getattr(core_props, prop_name, None)
                return date_obj.isoformat() if date_obj else ''
            except (AttributeError, TypeError):
                return ''

        logger.info("📦 Extracting metadata...")

        # Extraer autor y validar inmediatamente
        raw_author = safe_get_prop('author')
        logger.info(f"Autor extraído de core_properties: {raw_author}")

        # Validar si el autor de core_properties es válido (no es "autor" o valores genéricos)
        doc_author = raw_author
        doc_author_source = 'core_properties' if raw_author else 'unknown'

        if raw_author:
            # Valores genéricos que no son nombres reales de autores
            invalid_author_values = {
                'autor', 'author', 'unknown', 'desconocido', 'no especificado',
                'sin autor', 'sin nombre', 'n/a', 'n/a', 'n/a', 'n/a', 'n/a',
                'user', 'usuario', 'default', 'plantilla', 'template'
            }

            # Normalizar el autor para validación
            author_lower = raw_author.strip().lower()

            logger.info(f"Autor normalizado: {author_lower}")

            # Verificar si el autor es un valor genérico inválido
            if author_lower in invalid_author_values or not author_lower:
                logger.info(f"⚠️  Autor inválido encontrado en metadata: '{raw_author}'")
                doc_author = ''  # Limpiar el autor inválido
                doc_author_source = 'unknown'

        # --- Fallback: intentar extraer el autor desde tablas (incluyendo nested) ---
        if not doc_author:
            logger.info("No se encontró autor en core_properties, intentando extraer desde tablas")
            table_author = self._extract_author_from_tables(doc)
            if table_author:
                doc_author = table_author
                doc_author_source = 'table'

        # Crear diccionario de metadata con autor validado
        metadata = {
            # ========================================================================
            # PROPIEDADES DEL DOCUMENTO (Core Properties)
            # ========================================================================
            'doc_author': doc_author,
            'doc_author_source': doc_author_source,
            'doc_title': safe_get_prop('title') or file_path.stem,
            'doc_subject': safe_get_prop('subject'),
            'doc_keywords': safe_get_prop('keywords'),
            'doc_comments': safe_get_prop('comments'),
            'doc_category': safe_get_prop('category'),
            'doc_company': safe_get_prop('company'),
            'doc_last_modified_by': safe_get_prop('last_modified_by'),
            'doc_revision': safe_get_prop('revision', '1'),

            # ========================================================================
            # FECHAS (ISO 8601 para fácil filtrado en Qdrant)
            # ========================================================================
            'doc_created_date': safe_get_date('created'),
            'doc_modified_date': safe_get_date('modified'),

            # ========================================================================
            # INFORMACIÓN DEL ARCHIVO (para filtros en Qdrant)
            # ========================================================================
            'file_name': file_path.name,
            'file_extension': file_path.suffix,
            'file_size_bytes': file_stats.st_size,
            'file_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',

            # ========================================================================
            # ESTADÍSTICAS DEL DOCUMENTO
            # ========================================================================
            'total_paragraphs': total_paragraphs,
            'total_tables': total_tables,
            'language': detected_language,

            # METADATOS DE PROCESAMIENTO
            'indexed_at': get_current_utc_iso(),
            'processing_version': '2.5',
            'loader_type': 'WordLoaderUniversal',
            'parallel_processing_enabled': self.enable_parallel,
        }

        logger.debug(
            f"📋 Extracted metadata: title={metadata['doc_title']}, "
            f"author={metadata['doc_author']} (source={metadata.get('doc_author_source','unknown')}), "
            f"lang={metadata['language']}"
        )
        return metadata

    # -------------------------------------------------------------------------
    # Enumeración de TODAS las tablas (incluye nested tables dentro de celdas)
    # -------------------------------------------------------------------------
    def _enumerate_all_tables(self, doc: Document) -> List[Table]:
        """
        Retorna una lista con todas las tablas del documento, incluyendo anidadas
        dentro de celdas (buscadas vía XPath). Usa el parent de la celda para
        construir correctamente el objeto Table de python-docx.
        """
        tables: List[Table] = list(doc.tables)  # tablas de primer nivel
        try:
            for top in doc.tables:
                for row in top.rows:
                    for cell in row.cells:
                        nested_tables_xml = cell._element.xpath('.//w:tbl')
                        for nt_idx, tbl_xml in enumerate(nested_tables_xml):
                            try:
                                nested_table = Table(tbl_xml, cell._parent)
                                tables.append(nested_table)
                            except Exception as e:
                                logger.debug(f"Error building nested Table object: {e}")
        except Exception as e:
            logger.debug(f"Error enumerating all tables: {e}")
        return tables
    # -------------------------------------------------------------------------
    # Fallback para extraer autor desde tablas (Histórico de Versiones / Autorizaciones)
    # -------------------------------------------------------------------------
    def _extract_author_from_tables(self, doc: Document) -> Optional[str]:
        """
        Busca en todas las tablas una columna 'Autor' (o variantes) y retorna
        el último valor no vacío (asumiendo que corresponde a la versión más reciente).
        Además contempla el caso de la tabla de 'Autorizaciones', tomando la fila
        'Preparación' -> columna 'Nombre'.
        """
        author_candidates: List[str] = []

        def normalize(text: str) -> str:
            return re.sub(r'\s+', ' ', (text or '')).strip().lower()

        def is_author_header(text: str) -> bool:
            t = normalize(text)
            return t in {
                'autor', 'author',
                #'responsable',  # por si la plantilla usa este término
            }

        # Recorre todas las tablas (nivel superior + anidadas)
        logger.info(f"🔍 Analizando {len(self._enumerate_all_tables(doc))} tablas en el documento")
        for table_idx, table in enumerate(self._enumerate_all_tables(doc)):
            logger.info(f"📊 Procesando tabla {table_idx + 1}: {len(table.rows)} filas, {len(table.columns) if table.rows else 0} columnas")

            if not table.rows:
                logger.info(f"  ⚠️  Tabla {table_idx + 1} no tiene filas, saltando")
                continue

            # Mostrar encabezados de la tabla
            headers = [normalize(cell.text) for cell in table.rows[0].cells] if table.rows else []
            logger.info(f"  📋 Encabezados de tabla {table_idx + 1}: {headers}")

            # 1) Caso general: columna 'Autor' explícita
            author_col_idx = None
            try:
                author_col_idx = next(i for i, h in enumerate(headers) if is_author_header(h))
                logger.info(f"  ✅ Encontrada columna 'Autor' en la tabla {table_idx + 1}: índice {author_col_idx}")
            except StopIteration:
                # 1.a) Heurística del "Histórico de Versiones"
                # cabeceras típicas: ['versión','fecha','resumen de los cambios','autor']
                try:
                    headers_set = set(headers)
                    if {'versión', 'version'}.intersection(headers_set) \
                       and {'fecha'}.issubset(headers_set) \
                       and {'resumen de los cambios', 'resumen', 'summary'}.intersection(headers_set) \
                       and 'autor' in headers_set:
                        author_col_idx = headers.index('autor')
                        logger.info(f"  ✅ Encontrada columna 'Autor' por heurística en tabla {table_idx + 1}: índice {author_col_idx}")
                except Exception:
                    logger.info(f"  ❌ No se encontró columna 'Autor' en tabla {table_idx + 1}")
                    author_col_idx = None

            # Extraer autores por columna 'Autor'
            if author_col_idx is not None:
                logger.info(f"  📝 Extrayendo autores de la columna 'Autor' (índice {author_col_idx}) en tabla {table_idx + 1}")
                for row_idx, row in enumerate(list(table.rows)[1:], 2):  # Empezar desde la fila 2
                    # Mostrar contenido de todas las celdas de la fila
                    row_cells_content = [cell.text.strip() for cell in row.cells]
                    logger.info(f"    📄 Fila {row_idx}: {len(row.cells)} celdas -> {row_cells_content}")

                    if author_col_idx >= len(row.cells):
                        logger.info(f"    ⚠️  Fila {row_idx}: índice de columna {author_col_idx} fuera de rango ({len(row.cells)} celdas)")
                        continue
                    cell_text = (row.cells[author_col_idx].text or '').strip()
                    logger.info(f"    📝 Fila {row_idx}, columna {author_col_idx}: '{cell_text}'")
                    if cell_text:
                        logger.info(f"    ✅ Añadiendo autor encontrado: '{cell_text}'")
                        author_candidates.append(cell_text)
                # seguir a siguiente tabla (pero igual luego evaluamos 'Autorizaciones')
                # no 'continue' aquí para permitir detectar ambos casos en la misma tabla

            # 2) Caso especial: tabla de 'Autorizaciones'
            # estructura típica: cabecera ['','nombre','fecha'] y filas:
            # 'preparación' -> nombre del autor
            try:
                logger.info(f"  🔍 Buscando tabla de 'Autorizaciones' en tabla {table_idx + 1}")
                for row_idx, row in enumerate(list(table.rows)[1:], 2):  # Empezar desde la fila 2
                    logger.info(f"    📄 Fila {row_idx}: buscando 'preparación' en primera columna")
                    if len(row.cells) < 2:
                        logger.info(f"    ⚠️  Fila {row_idx}: necesita al menos 2 celdas, tiene {len(row.cells)}")
                        continue
                    first_col = normalize(row.cells[0].text)
                    logger.info(f"    📝 Fila {row_idx}, columna 0: '{first_col}'")
                    if first_col in {'preparación', 'preparacion'}:
                        prep_name = (row.cells[1].text or '').strip()
                        logger.info(f"    ✅ Encontrada fila 'preparación' en tabla {table_idx + 1}, fila {row_idx}: '{prep_name}'")
                        if prep_name:
                            logger.info(f"    ✅ Añadiendo autor de preparación: '{prep_name}'")
                            author_candidates.append(prep_name)
            except Exception as e:
                logger.debug(f"Error parsing 'Autorizaciones' style table in table {table_idx + 1}: {e}")
        # Regresar el autor con mayor prioridad
        # Prioridad 1: Autor de "Preparación" (el más específico)
        # Prioridad 2: Autor de columna "Autor" (el más reciente)

        # Buscar primero si hay algún autor de "Preparación"
        for val in reversed(author_candidates):
            cleaned = re.sub(r'\s+', ' ', val).strip()
            if cleaned:
                # Verificar si este autor vino de una fila "Preparación"
                # Necesitamos rastrear de dónde vino cada autor
                # Por ahora, asumimos que si hay múltiples, el de preparación debería ser el correcto
                # Podemos identificarlo por el contexto o simplemente priorizarlo

                # Para identificar el autor de preparación, necesitamos modificar la lógica
                # Vamos a cambiar el enfoque: almacenar tuplas (autor, tipo)
                pass

        # Regresar todos los autores posibles (con distinct)
        if not author_candidates:
            logger.info("❌ No se encontró ningún autor válido en las tablas")
            return None

        # Limpiar y eliminar duplicados manteniendo el orden
        cleaned_authors = []
        seen = set()

        for val in author_candidates:
            cleaned = re.sub(r'\s+', ' ', val).strip()
            if cleaned and cleaned not in seen:
                cleaned_authors.append(cleaned)
                seen.add(cleaned)

        logger.info(f"✅ Autores encontrados (sin duplicados): {cleaned_authors}")

        # Retornar el primer autor (el más reciente/confiable)
        # Pero ahora tienes todos los autores disponibles si necesitas usarlos
        if cleaned_authors:
            logger.info(f"✅ Usando primer autor como autor final: '{cleaned_authors[0]}'")
            return cleaned_authors[0]

        return None

    def _detect_language(self, text: str) -> str:
        """Detecta el idioma del texto (heurística simple)"""
        if not text or len(text) < 10:
            return 'unknown'

        spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'los', 'del', 'se', 'las', 'es', 'por', 'para', 'con']
        english_words = ['the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'it', 'is', 'for', 'on', 'with']

        text_lower = text.lower()
        spanish_count = sum(1 for word in spanish_words if f' {word} ' in text_lower)
        english_count = sum(1 for word in english_words if f' {word} ' in text_lower)

        if spanish_count > english_count:
            return 'es'
        elif english_count > spanish_count:
            return 'en'
        return 'unknown'

    # =========================================================================
    # EXTRACCIÓN DE TABLAS
    # =========================================================================

    def _extract_table_content(self, table: Table, table_index: int, depth: int = 0) -> Tuple[str, Dict]:
        """
        Extrae contenido de tabla en formato legible + metadata.

        SOPORTA TABLAS ANIDADAS (nested tables dentro de celdas)

        Args:
            table: Tabla a procesar
            table_index: Índice de la tabla
            depth: Profundidad de anidación (para tablas dentro de tablas)

        Returns:
            tuple: (contenido_formateado, metadata_tabla)
        """
        table_text = []
        headers = []
        row_count = 0
        has_merged_cells = False
        has_nested_tables = False

        for idx, row in enumerate(table.rows):
            row_cells = []

            for cell in row.cells:
                # EXTRACCIÓN RECURSIVA DE CONTENIDO DE CELDA
                cell_content = self._extract_cell_content(cell, depth)
                row_cells.append(cell_content)

                # Detectar si la celda tiene tablas anidadas
                if '[NESTED_TABLE' in cell_content:
                    has_nested_tables = True

                # Detectar celdas fusionadas
                try:
                    if hasattr(cell._element, 'tcPr'):
                        tc_pr = cell._element.tcPr
                        if tc_pr is not None:
                            # Buscar gridSpan
                            grid_span = tc_pr.find(qn('w:gridSpan'))
                            if grid_span is not None:
                                has_merged_cells = True
                            # Buscar vMerge
                            v_merge = tc_pr.find(qn('w:vMerge'))
                            if v_merge is not None:
                                has_merged_cells = True
                except (AttributeError, TypeError) as e:
                    logger.debug(f"Error detecting merged cells: {e}")

            row_text = ' | '.join(row_cells)
            if row_text.strip():
                table_text.append(row_text)
                row_count += 1

                # Primera fila como headers
                if idx == 0:
                    headers = row_cells

        # Metadata de la tabla
        table_metadata = {
            'is_table': True,
            'table_index': table_index,
            'table_number': table_index + 1,
            'table_caption': f"Tabla {table_index + 1}",
            'column_headers': headers,
            'row_count': row_count - 1 if headers else row_count,
            'column_count': len(headers) if headers else 0,
            'has_merged_cells': has_merged_cells,
            'has_nested_tables': has_nested_tables,
            'nesting_depth': depth,
        }

        content = '\n'.join(table_text) if table_text else ''

        # Clean table content if cleaner is enabled
        if self.text_cleaner:
            content = self.text_cleaner.clean_table_content(content)

        if depth == 0:  # Solo contar tablas principales
            self.stats["tables"] += 1

        logger.debug(f"📊 Table {table_index + 1} (depth={depth}): {row_count} rows, "
                    f"{len(headers)} cols, nested={has_nested_tables}")

        return content, table_metadata

    def _extract_cell_content(self, cell: _Cell, depth: int = 0) -> str:
        """
        Extrae contenido completo de una celda, incluyendo:
        - Texto de párrafos
        - Tablas anidadas (recursivamente)
        - Formato especial

        MEJORADO: Ahora detecta tablas en cualquier posición de la celda

        Args:
            cell: Celda a procesar
            depth: Profundidad de anidación

        Returns:
            str: Contenido formateado de la celda
        """
        cell_parts = []
        nested_table_counter = 0

        try:
            # Método 1: Iterar sobre elementos directos de la celda
            for element in cell._element:
                # Es un párrafo
                if isinstance(element, CT_P):
                    para = Paragraph(element, cell._parent)
                    text = para.text.strip()
                    if text:
                        cell_parts.append(text)

                # Es una tabla anidada
                elif isinstance(element, CT_Tbl):
                    nested_table = Table(element, cell._parent)

                    # Extraer tabla anidada recursivamente
                    nested_content, nested_metadata = self._extract_table_content(
                        nested_table,
                        nested_table_counter,
                        depth=depth + 1
                    )

                    if nested_content:
                        # Formatear tabla anidada con indentación
                        indent = "  " * depth
                        nested_formatted = f"\n{indent}[NESTED_TABLE_L{depth+1}]\n"

                        # Indentar cada línea de la tabla anidada
                        for line in nested_content.split('\n'):
                            if line.strip():  # Solo líneas no vacías
                                nested_formatted += f"{indent}  {line}\n"

                        nested_formatted += f"{indent}[/NESTED_TABLE_L{depth+1}]"
                        cell_parts.append(nested_formatted)

                        nested_table_counter += 1

                        # Contar tabla anidada en stats
                        self.stats["nested_tables"] = self.stats.get("nested_tables", 0) + 1

                        logger.debug(f"  {'  ' * depth}└─ Found nested table at depth {depth+1}")

            # Método 2: Buscar tablas que puedan estar en estructuras complejas
            # (algunas veces las tablas están dentro de otros contenedores)
            if nested_table_counter == 0:
                # Buscar tablas usando XPath como fallback
                nested_tables_xml = cell._element.xpath('.//w:tbl')
                for idx, tbl_element in enumerate(nested_tables_xml):
                    try:
                        nested_table = Table(tbl_element, cell._parent)

                        nested_content, nested_metadata = self._extract_table_content(
                            nested_table,
                            idx,
                            depth=depth + 1
                        )

                        if nested_content:
                            indent = "  " * depth
                            nested_formatted = f"\n{indent}[NESTED_TABLE_L{depth+1}_ALT]\n"

                            for line in nested_content.split('\n'):
                                if line.strip():
                                    nested_formatted += f"{indent}  {line}\n"

                            nested_formatted += f"{indent}[/NESTED_TABLE_L{depth+1}_ALT]"
                            cell_parts.append(nested_formatted)

                            self.stats["nested_tables"] = self.stats.get("nested_tables", 0) + 1

                            logger.debug(f"  {'  ' * depth}└─ Found nested table via XPath at depth {depth+1}")
                    except Exception as e:
                        logger.debug(f"Error extracting nested table via XPath: {e}")

        except Exception as e:
            logger.error(f"Error extracting cell content at depth {depth}: {e}")
            # Fallback al método antiguo
            return cell.text.strip()

        # Unir todas las partes de la celda
        result = ' '.join(cell_parts) if cell_parts else ''

        # Log para debugging
        if depth <= 2 and result:  # Solo log primeros niveles
            preview = result[:100] + '...' if len(result) > 100 else result
            logger.debug(f"  {'  ' * depth}Cell content (depth={depth}): {preview}")

        return result

    # =========================================================================
    # EXTRACCIÓN DE FORMATO
    # =========================================================================

    def _extract_paragraph_formatting(self, para: Paragraph) -> Dict[str, Any]:
        """
        Extrae información de formato del párrafo.

        Returns:
            dict: Información de formato para metadata
        """
        has_bold = any(run.bold for run in para.runs if run.bold)
        has_italic = any(run.italic for run in para.runs if run.italic)
        has_underline = any(run.underline for run in para.runs if run.underline)

        # Detectar listas
        is_list = para.style.name.startswith('List')
        list_level = None
        if is_list:
            # Extraer nivel de lista (List Bullet, List Bullet 2, etc.)
            level_match = re.search(r'\d+', para.style.name)
            list_level = int(level_match.group()) if level_match else 1
            self.stats["lists"] += 1

        # Tamaño de fuente promedio
        font_sizes = []
        for run in para.runs:
            try:
                if run.font.size:
                    font_sizes.append(run.font.size.pt)
            except:
                pass

        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 11

        # Detectar colores
        has_colored_text = False
        for run in para.runs:
            try:
                if run.font.color and run.font.color.rgb:
                    has_colored_text = True
                    break
            except:
                pass

        return {
            'has_bold_text': has_bold,
            'has_italic_text': has_italic,
            'has_underline': has_underline,
            'has_colored_text': has_colored_text,
            'is_list_item': is_list,
            'list_level': list_level,
            'font_size_avg': int(avg_font_size),
        }

    # =========================================================================
    # DETECCIÓN DE CONTENIDO ESPECIAL
    # =========================================================================

    def _detect_content_types(self, text: str) -> Dict[str, bool]:
        """
        Detecta tipos de contenido especial en el texto.
        Útil para filtros avanzados en búsquedas.
        """
        return {
            'contains_numbers': bool(re.search(r'\d+', text)),
            'contains_dates': bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', text)),
            'contains_urls': bool(re.search(r'https?://\S+', text)),
            'contains_emails': bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)),
            'contains_money': bool(re.search(r'[$€£¥]\s*\d+|(\d+)\s*[$€£¥]', text)),
            'contains_percentages': bool(re.search(r'\d+\s*%', text)),
        }

    def _extract_hyperlinks(self, para: Paragraph) -> List[Dict[str, str]]:
        """
        Extrae hyperlinks de un párrafo.

        Returns:
            List de dicts con {text, url}
        """
        hyperlinks = []

        try:
            for hyperlink in para._element.xpath('.//w:hyperlink'):
                # Texto del hyperlink
                text_elements = hyperlink.xpath('.//w:t')
                text = ''.join([t.text for t in text_elements if t.text])

                # URL del hyperlink
                rId = hyperlink.get(qn('r:id'))
                if rId:
                    try:
                        url = para.part.rels[rId].target_ref
                        hyperlinks.append({
                            "text": text,
                            "url": url
                        })
                        self.stats["hyperlinks"] += 1
                    except:
                        pass
        except Exception as e:
            logger.debug(f"Error extracting hyperlinks: {e}")

        return hyperlinks

    def _extract_images(self, para: Paragraph) -> List[Dict[str, Any]]:
        """
        Extrae imágenes de un párrafo.

        Returns:
            List de dicts con información de imagen
        """
        images = []

        try:
            # Buscar imágenes en runs
            for run in para.runs:
                # Buscar elementos de imagen en el XML
                inline_shapes = run._element.xpath('.//a:blip')

                for blip in inline_shapes:
                    # Obtener rId de la imagen
                    rId = blip.get(qn('r:embed'))
                    if rId:
                        try:
                            # Obtener la imagen del documento
                            image_part = para.part.related_parts[rId]
                            image_bytes = image_part.blob

                            # Convertir a base64 (opcional)
                            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                            # Metadata de la imagen
                            image_info = {
                                'type': 'image',
                                'format': image_part.content_type,
                                'size_bytes': len(image_bytes),
                                'base64': image_base64[:100] + '...',
                                'rId': rId,
                            }

                            images.append(image_info)
                            self.stats["images"] += 1

                        except Exception as e:
                            logger.debug(f"Error extracting image {rId}: {e}")
        except Exception as e:
            logger.debug(f"Error in image extraction: {e}")

        return images

    def _extract_textboxes(self, para: Paragraph) -> List[str]:
        """
        Extrae texto de cuadros de texto (textboxes).

        Returns:
            List de textos de textboxes
        """
        textboxes = []

        try:
            # Buscar textboxes en el XML
            for textbox in para._element.xpath('.//w:txbxContent'):
                textbox_paras = textbox.xpath('.//w:t')
                textbox_text = ''.join([t.text for t in textbox_paras if t.text])

                if textbox_text.strip():
                    textboxes.append(textbox_text.strip())
                    self.stats["textboxes"] += 1

        except Exception as e:
            logger.debug(f"Error extracting textboxes: {e}")

        return textboxes

    def _detect_equations(self, para: Paragraph) -> bool:
        """
        Detecta si el párrafo contiene ecuaciones matemáticas.

        Returns:
            bool: True si contiene ecuaciones
        """
        try:
            # Buscar elementos de ecuación (OMML - Office MathML)
            math_elements = para._element.xpath('.//m:oMath')
            if math_elements:
                self.stats["equations"] += len(math_elements)
                return True
        except Exception as e:
            logger.debug(f"Error detecting equations: {e}")

        return False

    # =========================================================================
    # ITERACIÓN DE BLOQUES
    # =========================================================================

    def _iter_block_items(self, parent):
        """
        Itera sobre párrafos y tablas en orden de aparición.
        CRÍTICO para mantener el orden correcto del contenido.
        """
        if hasattr(parent, 'element'):
            parent_elm = parent.element.body
        else:
            parent_elm = parent

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # =========================================================================
    # EXTRACCIÓN DE SECCIONES
    # =========================================================================

    def extract_sections(self, doc: Document, doc_metadata: Dict) -> List[DocumentSection]:
        """
        Extrae secciones con metadata enriquecida para cada chunk.

        Compatible con FileProcessor para construcción de payloads Qdrant.

        Args:
            doc: Documento Word
            doc_metadata: Metadatos del documento completo

        Returns:
            List[DocumentSection]: Secciones con metadata completa
        """
        sections = []
        current_section = None
        current_content = []
        current_level = 0
        current_chunk_metadata = []
        pre_heading_content = []
        pre_heading_metadata = []

        # Contadores globales
        paragraph_index = 0
        char_position = 0
        table_counter = 0

        # Breadcrumb jerárquico
        breadcrumb = []
        heading_stack = [None] * 6  # H1-H6

        logger.debug("🔍 Starting section extraction with enhanced features")

        # Iterar sobre todos los elementos
        for block in self._iter_block_items(doc):

            if isinstance(block, Paragraph):
                para = block

                # ====================================================================
                # DETECTAR HEADINGS
                # ====================================================================
                if para.style.name.startswith('Heading'):
                    # Guardar contenido antes del primer heading
                    if current_section is None and pre_heading_content:
                        intro_content = '\n\n'.join(pre_heading_content).strip()
                        if intro_content:
                            # Clean intro content if cleaner enabled
                            if self.text_cleaner:
                                intro_content = self.text_cleaner.clean_text(intro_content)
                            sections.append(DocumentSection(
                                title="Introduction",
                                content=intro_content,
                                level=0,
                                metadata={
                                    # Metadata del documento
                                    **doc_metadata,
                                    # Metadata de la sección
                                    'section_title': 'Introduction',
                                    'heading_level': 0,
                                    'breadcrumb': ['Introduction'],
                                    'parent_section': None,
                                    # Metadata de chunks/párrafos
                                    'chunk_metadata': pre_heading_metadata,
                                    'paragraph_count': len(pre_heading_metadata),
                                }
                            ))
                        pre_heading_content = []
                        pre_heading_metadata = []

                    # Guardar sección anterior
                    if current_section:
                        section_content = '\n\n'.join(current_content).strip()
                        if section_content:
                            # Clean section content if cleaner enabled
                            if self.text_cleaner:
                                section_content = self.text_cleaner.clean_text(section_content)
                            sections.append(DocumentSection(
                                title=current_section,
                                content=section_content,
                                level=current_level,
                                metadata={
                                    # Metadata del documento
                                    **doc_metadata,
                                    # Metadata de la sección
                                    'section_title': current_section,
                                    'heading_level': current_level,
                                    'breadcrumb': breadcrumb.copy(),
                                    'parent_section': breadcrumb[-2] if len(breadcrumb) > 1 else None,
                                    # Metadata de chunks/párrafos
                                    'chunk_metadata': current_chunk_metadata,
                                    'paragraph_count': len(current_chunk_metadata),
                                }
                            ))

                    # Nueva sección
                    level_match = para.style.name.replace('Heading ', '').replace('Heading', '1')
                    try:
                        current_level = int(level_match)
                    except ValueError:
                        current_level = 1

                    current_section = para.text.strip()
                    # Clean section title if cleaner enabled
                    if self.text_cleaner:
                        current_section = self.text_cleaner.clean_text(current_section)
                    current_content = []
                    current_chunk_metadata = []

                    # Actualizar breadcrumb
                    heading_stack[current_level - 1] = current_section
                    # Limpiar niveles inferiores
                    for i in range(current_level, 6):
                        heading_stack[i] = None
                    # Reconstruir breadcrumb
                    breadcrumb = [h for h in heading_stack[:current_level] if h]

                    logger.debug(f"📌 Section: '{current_section}' (L{current_level})")

                else:
                    # PROCESAR PÁRRAFO NORMAL CON OPTIMIZACIONES PARA RAG
                    # ================================================================
                    text = para.text.strip()
                    if text:
                        # Clean text if cleaner enabled
                        if self.text_cleaner:
                            text = self.text_cleaner.clean_text(text)

                        # Extraer formato
                        formatting = self._extract_paragraph_formatting(para)

                        # Detectar contenido especial
                        content_types = self._detect_content_types(text)

                        # Extraer hyperlinks
                        hyperlinks = self._extract_hyperlinks(para)

                        # Extraer imágenes (filtradas para RAG)
                        images = self._extract_images(para)
                        # Filter out base64 data for RAG optimization
                        if images and self.text_cleaner:
                            images = [img for img in images if self.text_cleaner.should_include_special_content('image', '')]

                        # Extraer textboxes
                        textboxes = self._extract_textboxes(para)

                        # Detectar ecuaciones
                        has_equations = self._detect_equations(para)

                        # Extraer comentarios del párrafo
                        para_comments = self._extract_paragraph_comments(para, self._global_cache['comments'])

                        # Extraer referencias a notas
                        note_refs = self._extract_paragraph_notes(para)

                        # Extraer shapes
                        shapes = self._extract_shapes(para)

                        # Extraer gráficos del párrafo
                        para_chart = self._extract_chart_from_paragraph(para)

                        # Extraer bookmarks del párrafo
                        para_bookmarks = self._extract_paragraph_bookmarks(para)

                        # Metadata del párrafo/chunk OPTIMIZADA para RAG
                        chunk_meta = {
                            'paragraph_index': paragraph_index,
                            'char_start': char_position,
                            'char_end': char_position + len(text),
                            'word_count': len(text.split()),
                            'text_hash': hashlib.sha256(text.encode()).hexdigest(),
                            **formatting,
                            **content_types,
                            'is_table': False,

                            # Hyperlinks (útil para RAG)
                            'has_hyperlinks': len(hyperlinks) > 0,
                            'hyperlink_count': len(hyperlinks),
                            'hyperlinks': hyperlinks if hyperlinks else None,

                            # Imágenes (sin base64 para optimizar)
                            'has_images': len(images) > 0,
                            'image_count': len(images),
                            # Remove base64 data from images for payload optimization
                            'images': [{'type': img['type'], 'format': img.get('format', 'unknown'),
                                       'size_bytes': img.get('size_bytes', 0)}
                                      for img in images] if images else None,

                            # Textboxes (útil para contexto)
                            'has_textboxes': len(textboxes) > 0,
                            'textbox_count': len(textboxes),
                            'textboxes': textboxes if textboxes else None,

                            # Ecuaciones
                            'has_equations': has_equations,

                            # Comentarios (útil para contexto)
                            'has_comments': len(para_comments) > 0,
                            'comment_count': len(para_comments),
                            'comments': para_comments if para_comments else None,

                            # Notas
                            'has_footnotes': len(note_refs['footnote_refs']) > 0,
                            'footnote_refs': note_refs['footnote_refs'] if note_refs['footnote_refs'] else None,
                            'has_endnotes': len(note_refs['endnote_refs']) > 0,
                            'endnote_refs': note_refs['endnote_refs'] if note_refs['endnote_refs'] else None,

                            # Shapes (solo si tienen texto útil)
                            'has_shapes': len(shapes) > 0,
                            'shape_count': len(shapes),
                            'shapes': shapes if shapes else None,

                            # Gráficos
                            'has_chart': para_chart is not None,
                            'chart': para_chart,

                            # Bookmarks
                            'has_bookmarks': len(para_bookmarks) > 0,
                            'bookmark_count': len(para_bookmarks),
                            'bookmarks': list(para_bookmarks.keys()) if para_bookmarks else None,
                        }

                        # Agregar elementos especiales al contenido de forma OPTIMIZADA
                        full_text = text

                        # Add textboxes with compact formatting
                        if textboxes and self.text_cleaner:
                            for tb in textboxes:
                                if self.text_cleaner.should_include_special_content('textbox', tb):
                                    formatted_tb = self.text_cleaner.format_special_content('textbox', tb)
                                    if formatted_tb:
                                        full_text += f'\n{formatted_tb}'

                        # Add comments with compact formatting
                        if para_comments and self.text_cleaner:
                            for comment in para_comments:
                                comment_text = comment.get('text', '')
                                if self.text_cleaner.should_include_special_content('comment', comment_text):
                                    formatted_comment = self.text_cleaner.format_special_content('comment', comment_text)
                                    if formatted_comment:
                                        full_text += f'\n{formatted_comment}'

                        # Add shapes with compact formatting
                        if shapes and self.text_cleaner:
                            for shape in shapes:
                                shape_text = shape.get('text', '')
                                if shape.get('has_text') and self.text_cleaner.should_include_special_content('shape', shape_text):
                                    formatted_shape = self.text_cleaner.format_special_content('shape', shape_text)
                                    if formatted_shape:
                                        full_text += f'\n{formatted_shape}'

                        # Final cleaning of full_text
                        if self.text_cleaner:
                            full_text = self.text_cleaner.clean_text(full_text)

                        if current_section:
                            current_content.append(full_text)
                            current_chunk_metadata.append(chunk_meta)
                        else:
                            pre_heading_content.append(full_text)
                            pre_heading_metadata.append(chunk_meta)

                        char_position += len(full_text) + 1
                        paragraph_index += 1
                        self.stats["paragraphs"] += 1

            elif isinstance(block, Table):
                # ====================================================================
                # PROCESAR TABLA
                # ====================================================================
                table_content, table_metadata = self._extract_table_content(block, table_counter)

                if table_content:
                    # Formato con marcadores
                    table_formatted = f"[TABLA {table_counter + 1}]\n{table_content}\n[/TABLA]"

                    # Metadata completa de la tabla
                    chunk_meta = {
                        'paragraph_index': paragraph_index,
                        'char_start': char_position,
                        'char_end': char_position + len(table_formatted),
                        'word_count': len(table_formatted.split()),
                        'text_hash': hashlib.sha256(table_formatted.encode()).hexdigest(),
                        **table_metadata,
                        **self._detect_content_types(table_content),
                    }

                    if current_section:
                        current_content.append(table_formatted)
                        current_chunk_metadata.append(chunk_meta)
                    else:
                        pre_heading_content.append(table_formatted)
                        pre_heading_metadata.append(chunk_meta)

                    char_position += len(table_formatted) + 1
                    paragraph_index += 1
                    table_counter += 1

        # ========================================================================
        # GUARDAR CONTENIDO FINAL
        # ========================================================================

        # Contenido previo si no hay headings
        if current_section is None and pre_heading_content:
            intro_content = '\n\n'.join(pre_heading_content).strip()
            if intro_content:
                # Clean intro content if cleaner enabled
                if self.text_cleaner:
                    intro_content = self.text_cleaner.clean_text(intro_content)
                sections.append(DocumentSection(
                    title="Document Content",
                    content=intro_content,
                    level=0,
                    metadata={
                        **doc_metadata,
                        'section_title': 'Document Content',
                        'heading_level': 0,
                        'breadcrumb': ['Document Content'],
                        'parent_section': None,
                        'chunk_metadata': pre_heading_metadata,
                        'paragraph_count': len(pre_heading_metadata),
                    }
                ))

        # Última sección
        if current_section:
            section_content = '\n\n'.join(current_content).strip()
            if section_content:
                # Clean section content if cleaner enabled
                if self.text_cleaner:
                    section_content = self.text_cleaner.clean_text(section_content)
                sections.append(DocumentSection(
                    title=current_section,
                    content=section_content,
                    level=current_level,
                    metadata={
                        **doc_metadata,
                        'section_title': current_section,
                        'heading_level': current_level,
                        'breadcrumb': breadcrumb.copy(),
                        'parent_section': breadcrumb[-2] if len(breadcrumb) > 1 else None,
                        'chunk_metadata': current_chunk_metadata,
                        'paragraph_count': len(current_chunk_metadata),
                    }
                ))

        logger.info(f"✅ Extraction complete: {len(sections)} sections, "
                   f"{paragraph_index} paragraphs, {table_counter} tables")

        return sections

    def _generate_full_content(self, sections: List[DocumentSection]) -> str:
        """
        Genera el contenido completo del documento desde las secciones.
        Usado por FileProcessor para document_preview.
        """
        full_content = []
        for section in sections:
            # Agregar título de sección con formato markdown
            header_prefix = '#' * (section.level if section.level > 0 else 1)
            full_content.append(f"{header_prefix} {section.title}")
            full_content.append(section.content)
            full_content.append("")  # Línea en blanco entre secciones

        return '\n'.join(full_content).strip()

    def debug_table_structure(self, file_path: Path) -> Dict[str, Any]:
        """
        Método de debugging para analizar estructura de tablas.

        Útil para diagnosticar problemas con tablas anidadas.

        Args:
            file_path: Ruta al documento

        Returns:
            Dict con información detallada de estructura
        """
        try:
            doc = Document(file_path)
        except Exception as e:
            return {"error": str(e)}

        structure = {
            "total_tables": len(doc.tables),
            "tables": []
        }

        for table_idx, table in enumerate(doc.tables):
            table_info = {
                "index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "cells_with_nested_tables": []
            }

            # Analizar cada celda
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Buscar tablas anidadas en la celda
                    nested_tables = cell._element.xpath('.//w:tbl')

                    if nested_tables:
                        cell_info = {
                            "position": f"row_{row_idx}_col_{col_idx}",
                            "nested_table_count": len(nested_tables),
                            "cell_text_preview": cell.text[:100] if cell.text else ""
                        }
                        table_info["cells_with_nested_tables"].append(cell_info)

            structure["tables"].append(table_info)

        return structure


# =============================================================================
# ALIAS PARA COMPATIBILIDAD
# =============================================================================

# Mantener nombre original para compatibilidad con tu código
WordLoader = WordLoaderUniversal
